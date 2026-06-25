"""Generate the Forge/Fabric compatibility analysis Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Forge-Fabric兼容方案.docx"

doc = Document()

# ── Global style ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph("")  # spacer
    return table

def add_code(text):
    """Add a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ═══════════════════════════════════════════════════════════════
# Title
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading("Forge / Fabric 模组服务端完全兼容方案", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("MC 隧道控制器 (mc-tunnel) — 技术预研文档").italic = True
doc.add_paragraph(f"版本: 1.0  |  日期: 2026-06-25  |  状态: 草案")
doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
# 1. 背景与目标
# ═══════════════════════════════════════════════════════════════
doc.add_heading("1. 背景与目标", level=1)

doc.add_paragraph(
    "当前 mc-tunnel 系统基于 PaperMC 服务端构建，核心管理功能（玩家管理、白名单、控制台命令、"
    "OP 管理）依赖 RCON 协议与 Minecraft 服务端通信。PaperMC 兼容 Bukkit/Spigot 插件体系，"
    "但不支持 Forge/Fabric 模组生态。"
)
doc.add_paragraph(
    "本方案分析将系统扩展至 Forge、Fabric、NeoForge 等模组服务端所需的全部改动，"
    "作为未来开发的技术参考。核心目标：在不破坏现有 PaperMC 体验的前提下，"
    "让用户可自由选择服务端类型。"
)

# ═══════════════════════════════════════════════════════════════
# 2. 服务端类型全景对比
# ═══════════════════════════════════════════════════════════════
doc.add_heading("2. 服务端类型全景对比", level=1)

doc.add_paragraph(
    "下表列出 5 种主流 Minecraft 服务端类型及其关键差异："
)

add_table(
    ["维度", "Vanilla", "PaperMC", "Forge", "Fabric", "NeoForge"],
    [
        ["扩展机制", "无", "插件 (.jar)\nplugins/ 目录", "模组 (.jar)\nmods/ 目录",
         "模组 (.jar)\nmods/ 目录", "模组 (.jar)\nmods/ 目录"],
        ["客户端要求", "原版即可", "原版即可", "客户端需装相同模组",
         "客户端需装相同模组", "客户端需装相同模组"],
        ["内置 RCON", "✅ 原生支持", "✅ 原生支持", "❌ 不支持",
         "❌ 不支持", "❌ 不支持"],
        ["下载源", "Mojang API", "papermc.io API", "files.minecraftforge.net",
         "meta.fabricmc.net", "maven.neoforged.net"],
        ["启动方式", "java -jar server.jar", "java -jar paper-xxx.jar",
         "需先运行 installer\n然后 java @args.txt", "java -jar fabric-launch.jar",
         "需先运行 installer\n然后 java @args.txt"],
        ["版本号格式", "1.21", "1.21-130", "1.20.1-47.2.0", "1.20.1-0.16.5-1.20.1",
         "1.20.1-47.2.0"],
        ["server.properties", "✅ 标准", "✅ 标准 + 扩展", "✅ 标准", "✅ 标准", "✅ 标准"],
        ["社区规模", "官方", "最大插件社区", "最大模组社区", "轻量模组社区", "Forge 的分支"],
    ],
)

doc.add_heading("2.1 推荐扩展顺序", level=2)
doc.add_paragraph(
    "根据技术难度和用户需求，建议按以下顺序逐步支持："
)
doc.add_paragraph("Vanilla (原版)", style="List Number")
doc.add_paragraph("PaperMC (已完成)", style="List Number")
doc.add_paragraph("Fabric (轻量、API 友好、社区活跃)", style="List Number")
doc.add_paragraph("Forge (生态最大、安装流程复杂)", style="List Number")
doc.add_paragraph("NeoForge (Forge 的精神继承者，可选)", style="List Number")

# ═══════════════════════════════════════════════════════════════
# 3. RCON 兼容性 — 核心难点
# ═══════════════════════════════════════════════════════════════
doc.add_heading("3. RCON 兼容性问题 —— 核心难点", level=1)

doc.add_paragraph(
    "RCON（Remote Console）是 Minecraft 1.9+ 内置的远程管理协议，PaperMC 原生支持。"
    "但 Forge 和 Fabric 默认不包含 RCON 实现，这意味着当前系统中所有依赖 RCON 的功能"
    "在模组服务端上将全部失效。"
)

doc.add_heading("3.1 受影响的功能清单", level=2)
doc.add_paragraph(
    "以下功能全部通过 RCON 实现，在无 RCON 环境下不可用："
)
add_table(
    ["功能模块", "RCON 命令", "影响程度"],
    [
        ["玩家列表", "list", "🔴 核心功能"],
        ["踢出玩家", "kick <player>", "🔴 核心功能"],
        ["OP 管理", "op / deop <player>", "🟡 中等"],
        ["白名单管理", "whitelist add/remove/list", "🔴 核心功能"],
        ["控制台命令", "任意命令", "🔴 核心功能"],
        ["玩家详情 (坐标/世界)", "data get entity", "🟡 中等"],
        ["世界/难度/模式查询", "difficulty / gamemode", "🟡 中等"],
        ["在线时长追踪", "控制台日志解析", "🟢 可降级"],
    ],
)

doc.add_heading("3.2 解决方案", level=2)

doc.add_paragraph("方案 A：自动安装 RCON 模组（推荐）", style="List Bullet")
doc.add_paragraph(
    "在下载 Forge/Fabric 服务端后，自动下载对应的 RCON 模组并放入 mods/ 目录。"
    "首次启动后模组自动生效，RCON 功能恢复。"
)
add_table(
    ["服务端", "RCON 模组", "下载源", "备注"],
    [
        ["Fabric", "ServerRedirect", "Modrinth / CurseForge",
         "功能完整，与原生 RCON 兼容"],
        ["Forge", "RCON", "CurseForge",
         "社区维护，部分版本滞后"],
        ["Forge 1.20+", "Server Redirect (Forge)", "Modrinth",
         "Fabric ServerRedirect 的 Forge 移植"],
        ["NeoForge", "同 Forge 模组", "—", "一般兼容"],
    ],
)
doc.add_paragraph(
    "⚠️ 风险：模组版本需要精确匹配 MC 版本，且部分模组可能有停更风险。"
    "需要在下载前做版本兼容性检查。"
)

doc.add_paragraph("方案 B：fallback 到控制台日志解析（不推荐）", style="List Bullet")
doc.add_paragraph(
    "如果 RCON 模组不可用，可以降级到解析服务端 stdout 输出来获取信息。"
    "这种方式不可靠、无法发送命令、且不同服务端的输出格式不一致。仅作为最后手段。"
)

# ═══════════════════════════════════════════════════════════════
# 4. 启动机制差异
# ═══════════════════════════════════════════════════════════════
doc.add_heading("4. 启动机制差异", level=1)

doc.add_heading("4.1 当前 PaperMC 启动流程", level=2)
add_code("java -Xmx2G -Xms1G -jar paper-1.21-130.jar nogui")
doc.add_paragraph(
    "一个命令即可启动，jar 文件自带所有依赖库。ProcessManager 直接管理此进程。"
)

doc.add_heading("4.2 Forge 启动流程", level=2)
doc.add_paragraph("Forge 需要两步操作：")

doc.add_paragraph("步骤 1：运行 installer 生成启动脚本", style="List Number")
add_code("java -jar forge-1.20.1-47.2.0-installer.jar --installServer")
doc.add_paragraph(
    "这会生成 run.sh / run.bat 和一个 libraries/ 目录（含数百个依赖 jar）。"
    "注意：installer 是交互式的，需要找到无头 (headless) 安装方式。"
    "Forge 1.17+ 支持 --installServer 参数。"
)

doc.add_paragraph("步骤 2：使用生成的参数文件启动", style="List Number")
add_code("java @libraries/net/minecraftforge/forge/1.20.1-47.2.0/unix_args.txt nogui")
doc.add_paragraph(
    "@ 语法是 Java 9+ 的参数文件功能，文件内每行一个 JVM 参数。"
    "Windows 上用 win_args.txt。需要根据操作系统选择正确的文件。"
)

doc.add_heading("4.3 Fabric 启动流程", level=2)
doc.add_paragraph("Fabric 相对简单，但仍比 PaperMC 多一步：")

doc.add_paragraph("步骤 1：下载 fabric-server-launch.jar（约 1MB）", style="List Number")
add_code("# 从 meta.fabricmc.net API 获取下载链接")
doc.add_paragraph("步骤 2：下载原版 server.jar 放到同目录", style="List Number")
add_code("java -jar fabric-server-launch.jar nogui")
doc.add_paragraph(
    "Fabric 启动器自动加载 mods/ 目录下的模组。但仍需额外下载原版 server.jar。"
)

doc.add_heading("4.4 需要的代码改动", level=2)

add_table(
    ["改动点", "当前", "目标"],
    [
        ["ProcessManager._cmd", "单个 list[str]\n例如 ['java', '-jar', 'paper.jar']",
         "支持多阶段启动\n(installer → 启动)"],
        ["adapter.start_server()", "直接启动 jar", "根据服务端类型选择启动方式"],
        ["首次安装", "下载 jar 即完成",
         "Forge: 下载 installer → 执行安装 → 验证\nFabric: 下载 launcher + 原版 jar"],
        ["JVM 参数", "用户自定义 Xmx/Xms", "Forge 安装器自动生成参数文件\n用户参数需合并"],
        ["jar 路径检测", "server/paper-*.jar", "按类型检测不同文件名模式"],
    ],
)

# ═══════════════════════════════════════════════════════════════
# 5. 下载源与版本管理
# ═══════════════════════════════════════════════════════════════
doc.add_heading("5. 下载源与版本管理", level=1)

doc.add_heading("5.1 各下载源 API 对比", level=2)

add_table(
    ["项目", "PaperMC", "Vanilla", "Forge", "Fabric"],
    [
        ["API 地址", "api.papermc.io/v2", "piston-meta.mojang.com", "files.minecraftforge.net",
         "meta.fabricmc.net/v2"],
        ["版本列表", "GET /projects/paper", "GET /mc/game/version_manifest.json",
         "GET /net/minecraftforge/forge/promos_slim.json", "GET /versions/loader"],
        ["下载 URL", "构建号拼 URL", "version.json → server.jar URL",
         "Maven 路径拼接", "loader + installer 版本组合"],
        ["稳定性", "✅ 非常稳定", "✅ 官方、稳定", "⚠️ 偶有宕机", "✅ 非常稳定"],
        ["速率限制", "几乎无", "有", "有", "几乎无"],
    ],
)

doc.add_heading("5.2 版本号格式差异", level=2)
doc.add_paragraph(
    "不同服务端的版本号格式完全不同，需要统一的内部版本模型："
)

add_table(
    ["服务端类型", "版本号示例", "结构说明"],
    [
        ["PaperMC", "1.21-130", "MC版本-构建号"],
        ["Vanilla", "1.21", "仅 MC 版本"],
        ["Forge", "1.20.1-47.2.0", "MC版本-Forge版本"],
        ["Fabric", "1.20.1-0.16.5-1.20.1", "MC版本-Loader版本-Mappings版本"],
    ],
)

doc.add_paragraph(
    "建议新增 ServerVersion 数据类，统一存储 {type, mc_version, build/loader_version, download_url}。"
)

doc.add_heading("5.3 downloader.py 需要的改动", level=2)
doc.add_paragraph(
    "当前的 downloader.py 仅对接 PaperMC API。需要新增以下下载器："
)
doc.add_paragraph("VanillaDownloader — Mojang API，下载 server.jar", style="List Bullet")
doc.add_paragraph("ForgeDownloader — Maven + installer 执行", style="List Bullet")
doc.add_paragraph("FabricDownloader — meta.fabricmc.net + Mojang (原版 jar)", style="List Bullet")
doc.add_paragraph(
    "建议用策略模式重构：定义 ServerDownloader 抽象基类，各类型实现 download() 和 list_versions()。"
)

# ═══════════════════════════════════════════════════════════════
# 6. 配置管理差异
# ═══════════════════════════════════════════════════════════════
doc.add_heading("6. 配置管理差异", level=1)

doc.add_heading("6.1 server.properties — 基本通用", level=2)
doc.add_paragraph(
    "好消息是 server.properties 在所有服务端类型上完全通用。PaperMC 有一些额外属性"
    "（如 enable-bstats、paper-settings），Forge 和 Fabric 会忽略不认识的属性。"
    "现有的 ServerPropertiesGenerator 几乎不需要改动。"
)

doc.add_heading("6.2 各类型特有配置", level=2)

add_table(
    ["服务端类型", "特有配置文件", "说明"],
    [
        ["PaperMC", "bukkit.yml, spigot.yml, paper.yml, commands.yml",
         "插件相关配置，已存在于 server/ 目录"],
        ["Forge", "forge.cfg, config/ 目录",
         "模组生成的配置文件，自动创建"],
        ["Fabric", "fabric-server-launcher.properties",
         "仅少量配置，通常无需修改"],
        ["Vanilla", "无", "只有 server.properties + ops.json + whitelist.json"],
    ],
)

doc.add_heading("6.3 config.yaml 改动", level=2)
doc.add_paragraph(
    "Config dataclass 需要新增字段来识别服务端类型："
)
add_code(
    "# config/loader.py — Config.mc 段新增:\n"
    'mc:\n'
    '  version: "1.21"\n'
    '  server_type: "papermc"    # NEW: papermc | vanilla | forge | fabric | neoforge\n'
    '  forge_version: ""          # NEW: Forge 版本号（仅 server_type=forge 时有效）\n'
    '  fabric_loader_version: ""  # NEW: Fabric Loader 版本（仅 server_type=fabric 时有效）'
)

# ═══════════════════════════════════════════════════════════════
# 7. 插件/模组管理差异
# ═══════════════════════════════════════════════════════════════
doc.add_heading("7. 插件 / 模组管理差异", level=1)

doc.add_heading("7.1 目录结构", level=2)
add_table(
    ["服务端", "扩展目录", "文件类型", "元数据来源"],
    [
        ["PaperMC", "server/plugins/", ".jar", "jar 内 plugin.yml"],
        ["Forge", "server/mods/", ".jar", "jar 内 META-INF/mods.toml"],
        ["Fabric", "server/mods/", ".jar", "jar 内 fabric.mod.json"],
        ["Vanilla", "无", "—", "—"],
    ],
)

doc.add_heading("7.2 UI 自适应", level=2)
doc.add_paragraph(
    '管理面板的「插件管理」区域需要根据 server_type 动态切换标签和路径：'
)
doc.add_paragraph('PaperMC → 标签显示「插件管理」，操作 plugins/ 目录', style="List Bullet")
doc.add_paragraph('Forge/Fabric → 标签显示「模组管理」，操作 mods/ 目录', style="List Bullet")
doc.add_paragraph("Vanilla → 隐藏该区域", style="List Bullet")
doc.add_paragraph(
    "后端同样需要按类型选择不同的目录扫描器和元数据解析器。同名接口，多态实现。"
)

doc.add_heading("7.3 元数据解析", level=2)
doc.add_paragraph(
    "比 PaperMC 插件管理更进一步，需要在后端解析 jar 内的元数据文件，"
    "提取名称、版本、作者、描述、依赖等信息，而不是只列文件名。"
)
add_code(
    "# Python 示例: 读取 jar 内的插件元数据\n"
    "import zipfile, yaml\n"
    "with zipfile.ZipFile('plugin.jar') as z:\n"
    "    with z.open('plugin.yml') as f:\n"
    "        meta = yaml.safe_load(f)\n"
    "# For Fabric: 读取 fabric.mod.json (JSON)\n"
    "# For Forge:  读取 META-INF/mods.toml (TOML)"
)

# ═══════════════════════════════════════════════════════════════
# 8. 世界存档兼容性
# ═══════════════════════════════════════════════════════════════
doc.add_heading("8. 世界存档兼容性", level=1)

doc.add_heading("8.1 兼容性矩阵", level=2)

add_table(
    ["从 ↓ / 到 →", "Vanilla", "PaperMC", "Forge", "Fabric"],
    [
        ["Vanilla", "✅", "✅ 安全", "⚠️ 单向", "⚠️ 单向"],
        ["PaperMC", "✅ 可能丢插件方块", "✅", "⚠️ 单向", "⚠️ 单向"],
        ["Forge", "❌ 丢失模组方块", "❌ 丢失模组方块", "✅", "❌"],
        ["Fabric", "❌ 丢失模组方块", "❌ 丢失模组方块", "❌", "✅"],
    ],
)

doc.add_paragraph(
    "⚠️ 重要警告：切换到模组端后，世界里会包含模组添加的方块/实体。"
    "如果切换回 PaperMC 或 Vanilla，这些内容会全部丢失。"
    "系统需要在版本切换时检测世界兼容性并给出醒目警告。"
)

doc.add_heading("8.2 WorldManager 需要的改动", level=2)
doc.add_paragraph(
    "当前的 WorldManager 在 world 迁移时检查「隐藏目录」和「level.dat 存在性」。"
    "面向模组端需要新增："
)
doc.add_paragraph(
    "记录世界创建时使用的服务端类型到 world_meta.json", style="List Bullet"
)
doc.add_paragraph(
    "切换服务端类型时检查存量世界是否兼容", style="List Bullet"
)
doc.add_paragraph(
    "切换到不兼容服务端时阻止启动或要求创建新世界", style="List Bullet"
)

# ═══════════════════════════════════════════════════════════════
# 9. 代码改动范围预估
# ═══════════════════════════════════════════════════════════════
doc.add_heading("9. 代码改动范围预估", level=1)

doc.add_heading("9.1 新增模块", level=2)
add_table(
    ["文件", "职责", "预估行数"],
    [
        ["core/mcserver/downloaders/__init__.py", "下载器策略模式入口", "~30"],
        ["core/mcserver/downloaders/vanilla.py", "Vanilla 下载器", "~150"],
        ["core/mcserver/downloaders/forge.py", "Forge 下载器 + installer 执行", "~300"],
        ["core/mcserver/downloaders/fabric.py", "Fabric 下载器", "~200"],
        ["core/mcserver/server_types.py", "服务端类型枚举 + 工厂函数", "~80"],
        ["core/mcserver/rcon_setup.py", "RCON 模组自动安装", "~150"],
        ["api/plugins.py", "插件/模组管理 API（通用）", "~200"],
    ],
)

doc.add_heading("9.2 需修改的现有文件", level=2)
add_table(
    ["文件", "改动内容", "预估改动行数"],
    [
        ["config/loader.py", "新增 server_type / forge_version 等字段", "~50"],
        ["config/defaults.yaml", "新增默认值", "~10"],
        ["core/mcserver/adapter.py", "按服务端类型选择启动命令和路径", "~100"],
        ["core/mcserver/downloader.py", "重构为策略模式，保留 Paper 下载器", "~80"],
        ["core/mcserver/worlds.py", "世界兼容性检查", "~60"],
        ["core/procman/manager.py", "支持多阶段启动 (install → run)", "~40"],
        ["main.py", "启动流程按 server_type 分支", "~30"],
        ["api/server.py", "版本下载 API 支持多类型", "~50"],
        ["web/static/app.js", "版本下载弹窗加类型选择 + 插件/模组管理 UI", "~400"],
        ["web/templates/admin.html", "插件/模组管理 HTML 结构", "~200"],
        ["web/static/style.css", "新增 UI 样式", "~50"],
    ],
)

doc.add_heading("9.3 总量估算", level=2)
add_table(
    ["类别", "预估行数"],
    [
        ["新增 Python 后端代码", "~1,110 行"],
        ["修改现有 Python 代码", "~420 行"],
        ["新增前端代码 (JS + HTML + CSS)", "~650 行"],
        ["新增测试代码", "~300 行"],
        ["合计", "~2,480 行"],
    ],
)
doc.add_paragraph(
    "按每人天 200-300 行高质量代码估算，纯开发约 8-12 人天。"
    "加上调研（RCON 模组版本兼容性验证）、集成测试、文档，"
    "完整交付约 3-4 周。"
)

# ═══════════════════════════════════════════════════════════════
# 10. 实施建议
# ═══════════════════════════════════════════════════════════════
doc.add_heading("10. 实施建议与阶段性里程碑", level=1)

doc.add_heading("10.1 推荐分四阶段实施", level=2)

add_table(
    ["阶段", "内容", "里程碑", "依赖"],
    [
        ["阶段 0\n(当前)", "PaperMC 插件管理\n(plugins/ 目录 + UI)",
         "用户可上传/删除/管理插件", "无"],
        ["阶段 A\n(Vanilla)", "支持 Vanilla 服务端\n+ 下载源 + 版本切换",
         "可选择 PaperMC 或 Vanilla", "阶段 0"],
        ["阶段 B\n(Fabric)", "支持 Fabric + RCON 模组自动安装\n+ mods/ 模组管理",
         "三种类型可选，模组端功能正常", "阶段 A"],
        ["阶段 C\n(Forge)", "支持 Forge + installer 流程\n+ 多阶段启动",
         "四种类型全覆盖", "阶段 A\n(可选并行 B+C)"],
    ],
)

doc.add_heading("10.2 风险清单", level=2)

add_table(
    ["风险", "概率", "影响", "缓解措施"],
    [
        ["RCON 模组停止维护", "🟡 中", "🔴 高",
         "调研多个备选方案；接受 fallback 到控制台解析"],
        ["Forge API 不稳定/下载慢", "🟡 中", "🟡 中",
         "多镜像源 + 缓存 + 断点续传"],
        ["Forge installer 在 Windows 上行为异常", "🟡 中", "🟡 中",
         "在所有目标 OS 上充分测试 installer 流程"],
        ["模组端切换导致世界损坏", "🟢 低", "🔴 高",
         "切换前自动备份世界；醒目警告 UI"],
        ["版本号体系复杂导致混淆", "🟡 中", "🟢 低",
         "UI 上清晰展示类型+版本；自动过滤不兼容组合"],
        ["Fabric/Forge 启动慢导致健康检查误判", "🟢 低", "🟡 中",
         "按服务端类型调整启动超时阈值"],
    ],
)

doc.add_heading("10.3 不建议做的事", level=2)
doc.add_paragraph("混合核心 (Mohist / Magma / Arclight)：同时支持插件+模组，但稳定性差、社区小、API 混乱。不值得投入。", style="List Bullet")
doc.add_paragraph("自动从 CurseForge / Modrinth 搜索下载模组：API 复杂、版权问题、用户应自主选择模组来源。只做上传管理即可。", style="List Bullet")
doc.add_paragraph("热重载模组/插件：Paper 的 /reload 命令本身就不稳定，插件开发者也不推荐。统一要求重启。", style="List Bullet")

# ── Footer ────────────────────────────────────────────────────
doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— 文档结束 —")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──────────────────────────────────────────────────────
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"[OK] Document saved to: {OUTPUT}")
